import docx
import os

file_dir = 'word_documents'

list_of_files = os.listdir(file_dir)

#Create empty wor document to merge all paras from all docs in this doc

new_doc = docx. Document()

for file in list_of_files:

    #creating file path
    file_path = os.path.join(file_dir,file)

    #mentioing each filepath to be document object
    word_doc = docx.Document(file_path)

    paragraphs = word_doc.paragraphs

    for paragraph in paragraphs:
        new_doc._body._element.addnext(paragraph._element)

        #add page break after each document
        #new_doc.add_page_break()



#save the document

new_doc.save(os.path.join(file_dir,'new_doc.docx'))


print("Merging complete Successfully!")












