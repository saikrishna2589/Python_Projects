# importing docx library to read, edit and write Word (.docx) document files
import docx

# storing file paths of two Word documents into variables
docpath1 = 'files/panda11.docx'  # target document where we want to insert paragraph
docpath2 = 'files/panda2.docx'   # source document from where we copy paragraph

# create Document objects for both files using docx.Document()
doc1 = docx.Document(docpath1)  # this is our base document
doc2 = docx.Document(docpath2)  # this is the document we are copying from

# access the first paragraph (index 0) from doc2
# doc2.paragraphs returns list of Paragraph objects; we select first one
paragraph2 = doc2.paragraphs[0]

# if we access paragraph2.text --> we only get plain text without formatting
# that's why we work with the Paragraph object directly to keep formatting

# access all paragraphs from doc1 (returns a list of Paragraph objects)
paragraph1 = doc1.paragraphs

# now we decide where to insert the copied paragraph in doc1
# paragraph1[1] means we select the second paragraph in doc1 (index starts at 0)

# _element is the underlying XML structure of the paragraph which includes formatting
# addnext() inserts the paragraph2 XML element right after paragraph1[1]
paragraph1[1]._element.addnext(paragraph2._element)

# finally, save the updated document to a new file
doc1.save('files/amend_doc.docx')
